from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "entrega-backend-gastos-balance-front.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(25, 61, 122)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(54, 54, 54)
    run.bold = True


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    run = p.add_run(text if not bold_prefix else "")
    if not bold_prefix:
        run.text = text
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Arial"
styles["Heading 2"].font.name = "Arial"

add_title(doc, "Entrega Backend: Gastos Diarios Integrados al Balance Premium por Sede")
add_paragraph(doc, "Fecha: 2026-06-02")

add_heading(doc, "Objetivo", 1)
add_paragraph(doc, "Se implementó en backend el manejo de gastos diarios para que el balance de ventas refleje también los egresos de caja del día y permita que el cuadre de caja coincida con el dinero real disponible.")
add_paragraph(doc, "Adicionalmente, esta funcionalidad quedó restringida a un plan premium por sede dentro de la plataforma.")
add_bullets(doc, [
    "Compra de verduras o insumos pagados el mismo día.",
    "Pago de recibos o servicios.",
    "Cualquier otro gasto operativo de la sede.",
])

add_heading(doc, "Resultado Funcional", 1)
add_paragraph(doc, "Antes, el balance sólo contemplaba ventas, costo de producción, inventario y cantidad de ventas. No existía un registro de egresos de caja manuales, por lo que el cuadre podía diferir del efectivo real.")
add_paragraph(doc, "Ahora, se puede registrar un gasto diario por sede y cada gasto queda asociado a sede, descripción, valor, modo de pago, fecha y administrador que lo registró. El balance general y por sede incluyen esos gastos y calculan la caja esperada.")
add_paragraph(doc, "El módulo de gastos sólo está habilitado para sedes con plan PREMIUM activo y al día.")

add_heading(doc, "Regla Comercial Nueva", 1)
add_bullets(doc, [
    "La suscripción se maneja por sede.",
    "Cada sede ahora puede tener plan BASICO o PREMIUM.",
    "El módulo de gastos sólo puede usarse si la sede tiene suscripción configurada.",
    "La sede debe estar activa y no vencida.",
    "La sede debe tener plan PREMIUM para usar gastos.",
])

add_heading(doc, "Modelo Implementado", 1)
add_bullets(doc, [
    "Nueva entidad: GastoDiario",
    "Nueva tabla: gasto_diario",
    "Campos: id, sedeId, administradorId, descripcion, valor, fecha, modoPago",
    "Nuevo campo de suscripción por sede: plan",
    "Valores posibles del plan: BASICO y PREMIUM",
])

add_heading(doc, "Endpoints Nuevos", 1)
add_heading(doc, "1. Crear gasto", 2)
add_paragraph(doc, "POST /api/administrador/gastos")
add_paragraph(doc, "Header requerido: Authorization: Bearer <token>")
add_paragraph(doc, "Body:")
add_code_block(doc, [
    "{",
    '  "sedeId": 1,',
    '  "descripcion": "Compra de verduras",',
    '  "valor": 45000,',
    '  "modoPago": "EFECTIVO"',
    "}",
])
add_paragraph(doc, "Reglas:")
add_bullets(doc, [
    "sedeId obligatorio.",
    "descripcion obligatoria.",
    "valor obligatorio y mayor a 0.",
    "modoPago obligatorio.",
    "El administrador autenticado debe tener acceso a la sede.",
    "La sede debe tener plan PREMIUM activo y vigente.",
])
add_paragraph(doc, "Respuesta esperada:")
add_code_block(doc, [
    "{",
    '  "error": false,',
    '  "respuesta": {',
    '    "id": 10,',
    '    "sedeId": 1,',
    '    "sedeNombre": "Sede Centro",',
    '    "descripcion": "Compra de verduras",',
    '    "valor": 45000,',
    '    "modoPago": "EFECTIVO",',
    '    "fecha": "2026-06-02T14:32:10",',
    '    "administradorId": 7,',
    '    "administradorNombre": "Juan Perez"',
    '  }',
    "}",
])
add_paragraph(doc, "Error esperado si la sede no tiene premium:")
add_code_block(doc, [
    "{",
    '  "error": true,',
    '  "respuesta": "La sede debe tener un plan PREMIUM activo para usar el modulo de gastos"',
    "}",
])

add_heading(doc, "2. Listar gastos", 2)
add_paragraph(doc, "GET /api/administrador/gastos")
add_bullets(doc, [
    "Query params soportados: empresaNit, sedeId, desde, hasta.",
    "Formato de fecha: YYYY-MM-DD.",
    "Si no se envían fechas, devuelve los gastos del día actual.",
    "Si el usuario es administrador delegado, sólo verá gastos de las sedes permitidas.",
])
add_paragraph(doc, "Ejemplo:")
add_paragraph(doc, "GET /api/administrador/gastos?sedeId=1&desde=2026-06-02&hasta=2026-06-02")

add_heading(doc, "Endpoints Existentes Modificados", 1)
add_bullets(doc, [
    "GET /api/administrador/balance/general",
    "GET /api/administrador/balance/sedes",
])
add_paragraph(doc, "Ambos endpoints ahora incluyen campos adicionales de gastos y caja.")
add_paragraph(doc, "Los gastos sólo se reflejan para sedes con plan PREMIUM habilitado.")

add_heading(doc, "Nuevos Campos del Balance", 1)
add_bullets(doc, [
    "totalGastos",
    "gastosEfectivo",
    "gastosTransferencia",
    "cajaEsperada",
    "utilidadNeta",
])
add_paragraph(doc, "Ejemplo de response de balance general:")
add_code_block(doc, [
    "{",
    '  "totalVentas": 500000,',
    '  "costoProduccion": 200000,',
    '  "utilidadBruta": 300000,',
    '  "totalGastos": 50000,',
    '  "gastosEfectivo": 30000,',
    '  "gastosTransferencia": 20000,',
    '  "cajaEsperada": 170000,',
    '  "utilidadNeta": 250000,',
    '  "valorInventario": 800000,',
    '  "stockTotal": 120,',
    '  "cantidadVentas": 35,',
    '  "ventasEfectivo": 200000,',
    '  "ventasTransferencia": 300000',
    "}",
])

add_heading(doc, "Reglas de Negocio", 1)
add_bullets(doc, [
    "utilidadBruta = totalVentas - costoProduccion",
    "utilidadNeta = utilidadBruta - totalGastos",
    "cajaEsperada = ventasEfectivo - gastosEfectivo",
    "Sólo los gastos pagados en EFECTIVO afectan cajaEsperada.",
    "Los gastos pagados en TRANSFERENCIA afectan utilidad, pero no el efectivo físico en caja.",
    "Si la sede no tiene PREMIUM, el módulo de gastos queda deshabilitado.",
])

add_heading(doc, "Cambio en Suscripciones", 1)
add_paragraph(doc, "La configuración de suscripción por sede ahora debe incluir el plan.")
add_code_block(doc, [
    "{",
    '  "sedeId": 1,',
    '  "plan": "PREMIUM",',
    '  "tipoCobro": "MENSUAL",',
    '  "precioMensual": 50000,',
    '  "precioAnual": 500000,',
    '  "fechaInicioServicio": "2026-06-02",',
    '  "observacion": "Plan premium con modulo de gastos",',
    '  "activa": true',
    "}",
])
add_bullets(doc, [
    "La respuesta de suscripción ahora incluye plan.",
    "La respuesta de suscripción ahora incluye gastosHabilitados.",
])

add_heading(doc, "Qué Debe Hacer Front", 1)
add_bullets(doc, [
    "Crear formulario de gastos con sede, descripción, valor y modo de pago.",
    "Usar los valores EFECTIVO y TRANSFERENCIA en modoPago.",
    "Consumir POST /api/administrador/gastos para registrar gastos.",
    "Consumir GET /api/administrador/gastos para historial o tabla.",
    "Actualizar la vista de balance para mostrar gastos y cajaEsperada.",
    "Actualizar el cuadre de caja para usar cajaEsperada en lugar de ventasEfectivo.",
    "Refrescar balance e historial de gastos después de guardar.",
    "Validar antes de habilitar gastos si la sede tiene plan PREMIUM.",
    "Si no tiene PREMIUM, mostrar mensaje o CTA de mejora de plan.",
])

add_heading(doc, "Payload Sugerido para Front", 1)
add_code_block(doc, [
    "{",
    '  "sedeId": 1,',
    '  "descripcion": "Pago recibo de energia",',
    '  "valor": 120000,',
    '  "modoPago": "TRANSFERENCIA"',
    "}",
])

add_heading(doc, "Resumen Técnico de Cambios", 1)
add_bullets(doc, [
    "Nueva entidad GastoDiario.",
    "Nuevo enum PlanSuscripcionSede.",
    "Nuevo repositorio GastoDiarioRepository.",
    "Nuevo servicio GastoDiarioServicio.",
    "Nuevo servicio SuscripcionFeatureService.",
    "Nuevo controlador GastoDiarioController.",
    "Nueva tabla gasto_diario.",
    "Nuevo campo plan en suscripcion_sede.",
    "Ajuste de BalanceGeneralDTO y BalanceSedeDTO.",
    "Ajuste de BalanceServicioImpl y BalanceController.",
    "Ajuste de DTOs de suscripción de sede.",
    "Ajuste de pruebas unitarias del balance.",
])

add_heading(doc, "Validación", 1)
add_paragraph(doc, "Se ejecutaron pruebas del backend con ./gradlew test y el resultado fue exitoso.")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
